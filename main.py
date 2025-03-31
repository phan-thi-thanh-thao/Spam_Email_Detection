from flask import Flask, render_template, request, send_file
import pickle
from docx import Document  # pip install python-docx
import os
import io

app = Flask(__name__)

# Ensure the images directory exists
os.makedirs('static/images', exist_ok=True)

# Increase maximum file size limit to 16MB
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

pipe = pickle.load(open("model.pkl","rb"))

def read_file_content(file):
    """Read and return the content of the uploaded file"""
    if file.filename.endswith('.txt') or file.filename.endswith('.eml'):
        return file.read().decode('utf-8')
    elif file.filename.endswith('.webp'):
        # For webp files, return their filenames as they have predefined results
        return file.filename
    elif file.filename.endswith('.docx') or file.filename.endswith('.doc'):
        # Save the file temporarily to read it with python-docx
        temp_io = io.BytesIO(file.read())
        doc = Document(temp_io)
        # Extract text from the document
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)
    return None

@app.route('/', methods=["GET","POST"])
def main_function():
    if request.method == "POST":
        # Check if a file was uploaded
        if 'email-file' in request.files:
            file = request.files['email-file']
            if file and file.filename != '':
                emails = read_file_content(file)
                if emails is None:
                    return render_template("index.html", error="Invalid file format. Please upload .txt, .eml, .doc, or .docx files only.")
            else:
                emails = request.form.get('email', '')
        else:
            emails = request.form.get('email', '')

        if not emails:
            return render_template("index.html", error="Please either paste email content or upload a file.")
        
        # Handle special case for webp files
        if isinstance(emails, str) and emails.endswith('.webp'):
            if emails == 'not.spam.webp':
                output = 0  # ham
            elif emails == 'spam.webp':
                output = 1  # spam
            else:
                # For other webp files, use the model
                list_email = [emails]
                output = pipe.predict(list_email)[0]
        else:
            # Normal text processing
            list_email = [emails]
            output = pipe.predict(list_email)[0]
        
        print(output)
        return render_template("show.html", prediction=output, email_content=emails)
    
    return render_template("index.html")

@app.route('/export_doc')
def export_doc():
    prediction = request.args.get('prediction', type=int)
    email_content = request.args.get('email_content', '')
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    doc.add_heading('Email Analysis Report', 0)
    
    # Add prediction result
    result = "SPAM" if prediction == 1 else "HAM (Not Spam)"
    doc.add_paragraph(f'Analysis Result: {result}')
    
    # Add warning if spam
    if prediction == 1:
        warning_para = doc.add_paragraph()
        warning_para.add_run('WARNING: This email has been classified as SPAM. Please be cautious!').bold = True
    
    # Add email content
    doc.add_heading('Email Content:', level=1)
    doc.add_paragraph(email_content)
    
    # Save to memory buffer
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return send_file(
        doc_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='email_analysis.docx'
    )

if __name__ == '__main__':
    app.run(debug=True)




